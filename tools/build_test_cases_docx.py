"""
将 CASES（与 build_test_cases.py 共用）输出为 Word 文档：
  -> D:\\zhubo-life-assisstant\\TEST_CASES.docx

文档结构：
  封面信息 → 测试准备 → 各模块用例（每条带前置/步骤/预期/状态/备注 5 行表格）→ 失败反馈模板
"""
import os
import sys
from collections import OrderedDict
from datetime import datetime

# 复用同目录下的 CASES
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_test_cases import CASES  # type: ignore

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "TEST_CASES.docx")


# ---------- 样式 / 工具 -----------------------------------------------------

def set_cell_bg(cell, hex_color: str) -> None:
    """给单元格加底色（hex 不含 #）"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        borders.append(b)
    tc_pr.append(borders)


def add_para(doc, text: str, *, bold=False, size: int = 11, color: str | None = None,
             align=WD_ALIGN_PARAGRAPH.LEFT, after_pt: int = 4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after_pt)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


# ---------- 内容渲染 --------------------------------------------------------

PRIO_COLOR = {"高": "FF3B30", "中": "FF9500", "低": "8E8E93"}

MODULE_DESC = OrderedDict([
    ("认证", "登录、退出、密码、会话存储等"),
    ("用户管理", "管理员对系统账号的增删改查（仅 admin 可见）"),
    ("主播", "主播档案的全生命周期"),
    ("节点", "成长轨迹节点 + 文件上传"),
    ("材料库", "目录树浏览 / 文件预览 / 删除 / 孤儿目录清理"),
    ("直播", "直播记录录入与统计"),
    ("培训", "培训计划、报名、签到、效果分析"),
    ("薪资", "薪资配置、预览、生成、发放"),
    ("合同", "签订、续签、终止、到期预警"),
    ("资产", "直播间、设备、借出、归还、逾期"),
    ("离职", "离职单提交与状态流转"),
    ("日志", "操作日志查阅"),
    ("看板", "经营数据汇总"),
    ("权限", "三种角色的权限矩阵"),
    ("端到端", "跨模块端到端流程"),
    ("安全", "边界与安全测试"),
    ("体验", "兼容、性能、外观"),
])


def stat_summary() -> dict:
    by_module: dict[str, int] = {}
    by_prio: dict[str, int] = {}
    core: list[tuple[str, str]] = []
    for tc_id, mod, title, prio, core_flag, *_ in CASES:
        by_module[mod] = by_module.get(mod, 0) + 1
        by_prio[prio] = by_prio.get(prio, 0) + 1
        if core_flag == "是":
            core.append((tc_id, title))
    return {"by_module": by_module, "by_prio": by_prio, "core": core, "total": len(CASES)}


def render_cover(doc: Document, stats: dict) -> None:
    title = doc.add_heading("主播全生命周期管理系统 — 手动测试用例", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, f"用例总数：{stats['total']} 条     生成时间：{datetime.now():%Y-%m-%d %H:%M}",
             align=WD_ALIGN_PARAGRAPH.CENTER, color="6E6E73", size=10, after_pt=12)

    # 测试环境
    doc.add_heading("测试准备", level=1)
    add_para(doc, "环境", bold=True, size=12)
    add_para(doc, "操作系统：Windows 10 / 11；Python 3.11；Node 18+", size=10)

    add_para(doc, "启动后端 / 前端", bold=True, size=12)
    code1 = doc.add_paragraph()
    code1.paragraph_format.left_indent = Cm(0.5)
    code1.add_run(
        "cd D:\\zhubo-life-assisstant\\backend\n"
        ".\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8000\n\n"
        "cd D:\\zhubo-life-assisstant\\frontend\n"
        "npm run dev"
    ).font.name = "Consolas"

    add_para(doc, "默认账号", bold=True, size=12)
    accounts = [
        ("用户名", "密码", "角色", "权限"),
        ("admin", "admin123", "管理员", "全部"),
        ("operator", "operator123", "运营", "业务读写；无薪资写 / 合同终止 / 用户管理"),
        ("finance", "finance123", "财务", "薪资读写；业务只读"),
    ]
    t = doc.add_table(rows=len(accounts), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(accounts):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = val
            if i == 0:
                set_cell_bg(cell, "F2F2F7")
                cell.paragraphs[0].runs[0].bold = True
            set_cell_borders(cell)
    doc.add_paragraph()

    add_para(doc, "数据基线", bold=True, size=12)
    add_para(doc, "每轮回归前先停后端，运行 tools\\reset_db.py 重置数据库与文件目录，"
                  "再启动后端，自动重建表与三个默认账号。", size=10)

    # 用例分布
    doc.add_heading("用例分布", level=1)
    add_para(doc, f"按优先级：高 {stats['by_prio'].get('高',0)} 条 / 中 {stats['by_prio'].get('中',0)} 条 / 低 {stats['by_prio'].get('低',0)} 条",
             size=11)

    add_para(doc, "按模块：", size=11, after_pt=2)
    mt = doc.add_table(rows=1, cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    head = mt.rows[0].cells
    for j, h in enumerate(["模块", "数量", "模块", "数量"]):
        head[j].text = h
        set_cell_bg(head[j], "F2F2F7")
        head[j].paragraphs[0].runs[0].bold = True
        set_cell_borders(head[j])
    items = list(stats["by_module"].items())
    for i in range(0, len(items), 2):
        row = mt.add_row().cells
        row[0].text = items[i][0]
        row[1].text = str(items[i][1])
        if i + 1 < len(items):
            row[2].text = items[i + 1][0]
            row[3].text = str(items[i + 1][1])
        for c in row:
            set_cell_borders(c)
    doc.add_paragraph()

    # 核心需求清单
    doc.add_heading("核心需求验收清单", level=1)
    add_para(doc, "本次重点验收的用例（必须全部通过）：", size=11)
    for tc_id, title in stats["core"]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{tc_id}  {title}")
        run.bold = True

    doc.add_page_break()

    # 测试约定
    doc.add_heading("测试约定", level=1)
    convs = [
        "状态填写：通过 / 失败 / 阻塞 / 跳过 / 待执行（默认）。",
        "失败用例请填写备注，记录现象、报错与定位线索。",
        "建议按当前文档章节顺序执行，前置数据后置依赖明确。",
        "如需重置数据，停后端 → 运行 tools\\reset_db.py → 重启后端。",
        "三种角色可同时多窗口/无痕窗口登录，SessionStorage 互不干扰。",
    ]
    for c in convs:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_page_break()


def render_case(doc: Document, tc) -> None:
    tc_id, mod, title, prio, core_flag, pre, steps, expect = tc

    # 用例标题
    h = doc.add_heading(level=2)
    r = h.add_run(f"{tc_id}  {title}")

    # 元信息小行
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    runs = [
        ("优先级 ", "6E6E73"),
        (prio, PRIO_COLOR.get(prio, "1C1C1E")),
        ("    模块 ", "6E6E73"),
        (mod, "1C1C1E"),
    ]
    for text, color in runs:
        rr = p.add_run(text)
        rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor.from_string(color)
        if color == PRIO_COLOR.get(prio, "1C1C1E") and text == prio:
            rr.bold = True
    if core_flag == "是":
        rr = p.add_run("    [核心需求]")
        rr.font.size = Pt(10)
        rr.bold = True
        rr.font.color.rgb = RGBColor.from_string("FF3B30")

    # 5 行 2 列表
    rows = [
        ("前置条件", pre or "-"),
        ("测试步骤", steps or "-"),
        ("预期结果", expect or "-"),
        ("状态", "□ 通过    □ 失败    □ 阻塞    □ 跳过"),
        ("备注 / 失败截图", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(13.0)
    for i, (k, v) in enumerate(rows):
        row = table.rows[i]
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(13.0)

        c0 = row.cells[0]
        c0.text = k
        c0.paragraphs[0].runs[0].bold = True
        c0.paragraphs[0].runs[0].font.size = Pt(10)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_bg(c0, "F2F2F7")
        set_cell_borders(c0)

        c1 = row.cells[1]
        # 多行步骤：按 \n 拆段落，避免单段太长
        text_lines = (v or "-").split("\n")
        c1.text = text_lines[0]
        for extra in text_lines[1:]:
            c1.add_paragraph(extra)
        for para in c1.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_borders(c1)

    doc.add_paragraph()  # 分隔


def render_modules(doc: Document) -> None:
    # 按 MODULE_DESC 顺序遍历，未列入的模块按原顺序追加
    seen = set()
    grouped: "OrderedDict[str, list]" = OrderedDict()
    for tc in CASES:
        mod = tc[1]
        grouped.setdefault(mod, []).append(tc)
    ordered_keys = [k for k in MODULE_DESC.keys() if k in grouped]
    for k in grouped.keys():
        if k not in ordered_keys:
            ordered_keys.append(k)

    for mod in ordered_keys:
        if mod in seen:
            continue
        seen.add(mod)
        doc.add_heading(f"{mod} 模块", level=1)
        desc = MODULE_DESC.get(mod)
        if desc:
            add_para(doc, desc, color="6E6E73", size=10, after_pt=8)

        for tc in grouped[mod]:
            render_case(doc, tc)

        doc.add_page_break()


def render_feedback(doc: Document) -> None:
    doc.add_heading("失败用例反馈模板", level=1)
    add_para(doc, "把下列信息复制给开发，可以最快定位问题：", size=11)
    template = (
        "用例编号：TC-XXX-NN\n"
        "浏览器：Chrome 130\n"
        "登录角色：admin / operator / finance\n"
        "现象：点击 X 后页面 Y，应该 Z\n"
        "后端控制台报错：粘贴最后 20 行 traceback\n"
        "前端 Console 报错：粘贴红色那段\n"
        "网络请求：DevTools → Network 该接口的 Request / Response\n"
        "（截图可直接拖入对话）"
    )
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(template)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


# ---------- main -----------------------------------------------------------

def main() -> None:
    doc = Document()

    # 全局默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    rpr.append(rfonts)

    stats = stat_summary()
    render_cover(doc, stats)
    render_modules(doc)
    render_feedback(doc)

    doc.save(OUT)
    print(f"OK -> {OUT}; cases = {stats['total']}")


if __name__ == "__main__":
    main()
